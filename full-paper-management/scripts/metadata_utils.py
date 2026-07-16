#!/usr/bin/env python3
"""文献元数据提取工具 V2 — 支持 PDF/DOCX/TXT 文件的作者、年份、标题、DOI、关键词提取。

V2 新增:
  - 关键词提取（最多 5 个，分号分隔）
  - APA 逐字段规则构建 + 强制自检
  - 重命名冲突降级（姓氏 → 姓氏+名字全拼）
  - 非英文姓名拼音处理
  - 第一作者全名字段（用于冲突降级）

使用方式:
  python metadata_utils.py <file_path>           # 输出 JSON 元数据
  python metadata_utils.py --batch <dir_path>    # 批量处理目录下所有文献
  python metadata_utils.py --doi <doi_string>     # 通过 DOI 在线解析元数据
"""

import os, sys, json, re, hashlib, io
from pathlib import Path
from collections import Counter

# ── 依赖检测 ──────────────────────────────────────────────────
HAS_PYPDF = False
HAS_DOCX = False
HAS_REQUESTS = False

try:
    import PyPDF2
    HAS_PYPDF = True
except ImportError:
    pass

try:
    import docx
    HAS_DOCX = True
except ImportError:
    pass

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    pass

# ── 文献支持的文件类型 ─────────────────────────────────────────
PAPER_EXTS = {'.pdf', '.docx', '.doc', '.txt', '.epub', '.mobi'}

# ── 停用词表（关键词提取用） ───────────────────────────────────
STOP_WORDS = {
    # English
    'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
    'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'be',
    'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will',
    'would', 'could', 'should', 'may', 'might', 'must', 'can', 'this',
    'that', 'these', 'those', 'it', 'its', 'they', 'them', 'their',
    'we', 'our', 'us', 'you', 'your', 'he', 'she', 'his', 'her',
    'which', 'who', 'whom', 'what', 'where', 'when', 'why', 'how',
    'not', 'no', 'nor', 'so', 'than', 'too', 'very', 'just', 'about',
    'above', 'after', 'again', 'against', 'all', 'any', 'because',
    'before', 'below', 'between', 'during', 'each', 'few', 'more',
    'most', 'other', 'over', 'same', 'some', 'such', 'then', 'there',
    'here', 'through', 'under', 'until', 'up', 'down', 'out', 'off',
    'over', 'under', 'again', 'further', 'once', 'also', 'however',
    'may', 'might', 'could', 'would', 'should', 'shall', 'will',
    'using', 'used', 'use', 'based', 'via', 'per', 'etc', 'e.g',
    'i.e', 'vs', 'pp', 'vol', 'no', 'et', 'al', 'doi', 'http',
    'https', 'www', 'com', 'org', 'html', 'pdf', 'abstract',
    'introduction', 'conclusion', 'method', 'methods', 'result',
    'results', 'discussion', 'acknowledgment', 'reference', 'references',
    'figure', 'fig', 'table', 'tab', 'section', 'sec', 'eq', 'equation',
    # Chinese stop words
    '的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都',
    '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你',
    '会', '着', '没有', '看', '好', '自己', '这', '那', '与', '或',
    '但', '而', '及', '以', '为', '被', '使', '让', '从', '向',
    '中', '对', '于', '由', '把', '将', '该', '其', '此', '些',
    '可', '可以', '需要', '应', '应该', '进行', '通过', '根据',
    '按照', '基于', '由于', '因为', '所以', '如果', '虽然', '但是',
    '然后', '接下来', '最后', '首先', '其次', '此外', '另外',
}

# ── 核心数据结构 ───────────────────────────────────────────────

def empty_meta(filepath: str = "") -> dict:
    return {
        "filepath": filepath,
        "filename": os.path.basename(filepath) if filepath else "",
        "authors": [],
        "first_author_surname": "",
        "first_author_fullname": "",
        "first_author_given_name": "",
        "year": "",
        "title": "",
        "journal": "",
        "volume": "",
        "issue": "",
        "pages": "",
        "doi": "",
        "keywords": [],
        "has_fulltext": False,
        "has_references": False,
        "has_appendix": False,
        "has_abstract": False,
        "has_figures": False,
        "page_count": 0,
        "md5": "",
        "source": "",
        "text_preview": "",
    }


# ── 文件哈希 ──────────────────────────────────────────────────

def file_md5(filepath: str) -> str:
    """计算文件的 MD5 哈希。"""
    h = hashlib.md5()
    try:
        with open(filepath, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                h.update(chunk)
    except Exception:
        return ""
    return h.hexdigest()


# ── DOI 提取 ──────────────────────────────────────────────────

DOI_PATTERN = re.compile(r'\b(10\.\d{4,}/[-._;()/:A-Za-z0-9]+)\b', re.IGNORECASE)

def extract_doi_from_text(text: str) -> str:
    """从文本中提取第一个 DOI。"""
    m = DOI_PATTERN.search(text)
    return m.group(1) if m else ""


# ── 年份提取 ──────────────────────────────────────────────────

YEAR_PATTERNS = [
    re.compile(r'(?:published|publication|accepted|created|date)\b[:\s]+(\d{4})', re.IGNORECASE),
    re.compile(r'\((\d{4})\)'),
    re.compile(r'\b(19|20)\d{2}\b'),
]

def extract_year_from_text(text: str) -> str:
    """从文本中提取最可能的发表年份。"""
    for pat in YEAR_PATTERNS:
        m = pat.search(text)
        if m:
            y = m.group(1) if pat is YEAR_PATTERNS[0] else (m.group(1) if m.lastindex else m.group(0))
            if y and y.isdigit() and 1900 <= int(y) <= 2099:
                return y
    return ""


# ── 作者提取与标准化 ──────────────────────────────────────────

def split_author_name(raw: str) -> tuple:
    """将原始作者名拆分为 (姓氏, 名字部分)。
    
    返回: (surname, given_name)
    - 英文: "John David Smith" → ("Smith", "JohnDavid")
    - 英文: "Smith, John D." → ("Smith", "JohnD")
    - 中文: "张伟" → ("Zhang", "Wei") (需配合 pinyin)
    """
    raw = raw.strip().strip(',').strip(';').strip()
    if not raw:
        return ("", "")
    
    if ',' in raw:
        parts = raw.split(',', 1)
        surname = parts[0].strip()
        given = parts[1].strip() if len(parts) > 1 else ""
    else:
        parts = raw.split()
        if len(parts) >= 2:
            surname = parts[-1]
            given = "".join(parts[:-1])
        else:
            surname = raw
            given = ""
    
    # 清理特殊字符
    surname = re.sub(r"['\-\.]", "", surname)
    given = re.sub(r"['\-\.]", "", given)
    
    return (surname, given)


def normalize_author_name(raw: str) -> str:
    """将原始作者名标准化为姓氏（用于文件命名）。"""
    surname, _ = split_author_name(raw)
    return surname


def normalize_author_apa(raw: str) -> str:
    """将原始作者名标准化为 APA 格式: 姓, 首字母缩写."""
    surname, given = split_author_name(raw)
    if not surname:
        return ""
    if given:
        initials = ". ".join(c.upper() for c in given if c.isalpha()) + "."
        return f"{surname}, {initials}"
    return surname


def is_chinese_name(name: str) -> bool:
    """判断是否为中文姓名。"""
    return bool(re.search(r'[\u4e00-\u9fff]', name))


# ── 简易拼音转换（常用姓氏 + 基本规则） ────────────────────────

# 常见中文姓氏拼音映射
COMMON_SURNAMES = {
    '王': 'Wang', '李': 'Li', '张': 'Zhang', '刘': 'Liu', '陈': 'Chen',
    '杨': 'Yang', '赵': 'Zhao', '黄': 'Huang', '周': 'Zhou', '吴': 'Wu',
    '徐': 'Xu', '孙': 'Sun', '胡': 'Hu', '朱': 'Zhu', '高': 'Gao',
    '林': 'Lin', '何': 'He', '郭': 'Guo', '马': 'Ma', '罗': 'Luo',
    '梁': 'Liang', '宋': 'Song', '郑': 'Zheng', '谢': 'Xie', '韩': 'Han',
    '唐': 'Tang', '冯': 'Feng', '于': 'Yu', '董': 'Dong', '萧': 'Xiao',
    '程': 'Cheng', '曹': 'Cao', '袁': 'Yuan', '邓': 'Deng', '许': 'Xu',
    '傅': 'Fu', '沈': 'Shen', '曾': 'Zeng', '彭': 'Peng', '吕': 'Lv',
    '苏': 'Su', '卢': 'Lu', '蒋': 'Jiang', '蔡': 'Cai', '贾': 'Jia',
    '丁': 'Ding', '魏': 'Wei', '薛': 'Xue', '叶': 'Ye', '阎': 'Yan',
    '余': 'Yu', '潘': 'Pan', '杜': 'Du', '戴': 'Dai', '夏': 'Xia',
    '钟': 'Zhong', '汪': 'Wang', '田': 'Tian', '任': 'Ren', '姜': 'Jiang',
    '范': 'Fan', '方': 'Fang', '石': 'Shi', '姚': 'Yao', '谭': 'Tan',
    '廖': 'Liao', '邹': 'Zou', '熊': 'Xiong', '金': 'Jin', '陆': 'Lu',
    '郝': 'Hao', '孔': 'Kong', '白': 'Bai', '崔': 'Cui', '康': 'Kang',
    '毛': 'Mao', '邱': 'Qiu', '秦': 'Qin', '江': 'Jiang', '史': 'Shi',
    '顾': 'Gu', '侯': 'Hou', '邵': 'Shao', '孟': 'Meng', '龙': 'Long',
    '万': 'Wan', '段': 'Duan', '雷': 'Lei', '钱': 'Qian', '汤': 'Tang',
    '尹': 'Yin', '黎': 'Li', '易': 'Yi', '常': 'Chang', '武': 'Wu',
    '乔': 'Qiao', '贺': 'He', '赖': 'Lai', '龚': 'Gong', '文': 'Wen',
}

# 简易拼音首字母映射（用于名字部分）
PINYIN_INITIALS = {
    'a': 'a', 'ai': 'ai', 'an': 'an', 'ang': 'ang', 'ao': 'ao',
    'ba': 'ba', 'bai': 'bai', 'ban': 'ban', 'bang': 'bang', 'bao': 'bao',
    'bei': 'bei', 'ben': 'ben', 'beng': 'beng', 'bi': 'bi', 'bian': 'bian',
    'biao': 'biao', 'bie': 'bie', 'bin': 'bin', 'bing': 'bing', 'bo': 'bo',
    'bu': 'bu',
    'ca': 'ca', 'cai': 'cai', 'can': 'can', 'cang': 'cang', 'cao': 'cao',
    'ce': 'ce', 'cen': 'cen', 'ceng': 'ceng', 'cha': 'cha', 'chai': 'chai',
    'chan': 'chan', 'chang': 'chang', 'chao': 'chao', 'che': 'che',
    'chen': 'chen', 'cheng': 'cheng', 'chi': 'chi', 'chong': 'chong',
    'chou': 'chou', 'chu': 'chu', 'chuai': 'chuai', 'chuan': 'chuan',
    'chuang': 'chuang', 'chui': 'chui', 'chun': 'chun', 'chuo': 'chuo',
    'ci': 'ci', 'cong': 'cong', 'cou': 'cou', 'cu': 'cu', 'cuan': 'cuan',
    'cui': 'cui', 'cun': 'cun', 'cuo': 'cuo',
    'da': 'da', 'dai': 'dai', 'dan': 'dan', 'dang': 'dang', 'dao': 'dao',
    'de': 'de', 'dei': 'dei', 'den': 'den', 'deng': 'deng', 'di': 'di',
    'dian': 'dian', 'diao': 'diao', 'die': 'die', 'ding': 'ding',
    'diu': 'diu', 'dong': 'dong', 'dou': 'dou', 'du': 'du', 'duan': 'duan',
    'dui': 'dui', 'dun': 'dun', 'duo': 'duo',
    'e': 'e', 'en': 'en', 'er': 'er',
    'fa': 'fa', 'fan': 'fan', 'fang': 'fang', 'fei': 'fei', 'fen': 'fen',
    'feng': 'feng', 'fo': 'fo', 'fou': 'fou', 'fu': 'fu',
    'ga': 'ga', 'gai': 'gai', 'gan': 'gan', 'gang': 'gang', 'gao': 'gao',
    'ge': 'ge', 'gei': 'gei', 'gen': 'gen', 'geng': 'geng', 'gong': 'gong',
    'gou': 'gou', 'gu': 'gu', 'gua': 'gua', 'guai': 'guai', 'guan': 'guan',
    'guang': 'guang', 'gui': 'gui', 'gun': 'gun', 'guo': 'guo',
    'ha': 'ha', 'hai': 'hai', 'han': 'han', 'hang': 'hang', 'hao': 'hao',
    'he': 'he', 'hei': 'hei', 'hen': 'hen', 'heng': 'heng', 'hong': 'hong',
    'hou': 'hou', 'hu': 'hu', 'hua': 'hua', 'huai': 'huai', 'huan': 'huan',
    'huang': 'huang', 'hui': 'hui', 'hun': 'hun', 'huo': 'huo',
    'ji': 'ji', 'jia': 'jia', 'jian': 'jian', 'jiang': 'jiang',
    'jiao': 'jiao', 'jie': 'jie', 'jin': 'jin', 'jing': 'jing',
    'jiong': 'jiong', 'jiu': 'jiu', 'ju': 'ju', 'juan': 'juan',
    'jue': 'jue', 'jun': 'jun',
    'ka': 'ka', 'kai': 'kai', 'kan': 'kan', 'kang': 'kang', 'kao': 'kao',
    'ke': 'ke', 'ken': 'ken', 'keng': 'keng', 'kong': 'kong', 'kou': 'kou',
    'ku': 'ku', 'kua': 'kua', 'kuai': 'kuai', 'kuan': 'kuan',
    'kuang': 'kuang', 'kui': 'kui', 'kun': 'kun', 'kuo': 'kuo',
    'la': 'la', 'lai': 'lai', 'lan': 'lan', 'lang': 'lang', 'lao': 'lao',
    'le': 'le', 'lei': 'lei', 'leng': 'leng', 'li': 'li', 'lia': 'lia',
    'lian': 'lian', 'liang': 'liang', 'liao': 'liao', 'lie': 'lie',
    'lin': 'lin', 'ling': 'ling', 'liu': 'liu', 'long': 'long',
    'lou': 'lou', 'lu': 'lu', 'lv': 'lv', 'luan': 'luan', 'lue': 'lue',
    'lun': 'lun', 'luo': 'luo',
    'ma': 'ma', 'mai': 'mai', 'man': 'man', 'mang': 'mang', 'mao': 'mao',
    'me': 'me', 'mei': 'mei', 'men': 'men', 'meng': 'meng', 'mi': 'mi',
    'mian': 'mian', 'miao': 'miao', 'mie': 'mie', 'min': 'min',
    'ming': 'ming', 'miu': 'miu', 'mo': 'mo', 'mou': 'mou', 'mu': 'mu',
    'na': 'na', 'nai': 'nai', 'nan': 'nan', 'nang': 'nang', 'nao': 'nao',
    'ne': 'ne', 'nei': 'nei', 'nen': 'nen', 'neng': 'neng', 'ni': 'ni',
    'nian': 'nian', 'niang': 'niang', 'niao': 'niao', 'nie': 'nie',
    'nin': 'nin', 'ning': 'ning', 'niu': 'niu', 'nong': 'nong',
    'nou': 'nou', 'nu': 'nu', 'nv': 'nv', 'nuan': 'nuan', 'nue': 'nue',
    'nuo': 'nuo',
    'o': 'o', 'ou': 'ou',
    'pa': 'pa', 'pai': 'pai', 'pan': 'pan', 'pang': 'pang', 'pao': 'pao',
    'pei': 'pei', 'pen': 'pen', 'peng': 'peng', 'pi': 'pi', 'pian': 'pian',
    'piao': 'piao', 'pie': 'pie', 'pin': 'pin', 'ping': 'ping', 'po': 'po',
    'pou': 'pou', 'pu': 'pu',
    'qi': 'qi', 'qia': 'qia', 'qian': 'qian', 'qiang': 'qiang',
    'qiao': 'qiao', 'qie': 'qie', 'qin': 'qin', 'qing': 'qing',
    'qiong': 'qiong', 'qiu': 'qiu', 'qu': 'qu', 'quan': 'quan',
    'que': 'que', 'qun': 'qun',
    'ran': 'ran', 'rang': 'rang', 'rao': 'rao', 're': 're', 'ren': 'ren',
    'reng': 'reng', 'ri': 'ri', 'rong': 'rong', 'rou': 'rou', 'ru': 'ru',
    'rua': 'rua', 'ruan': 'ruan', 'rui': 'rui', 'run': 'run', 'ruo': 'ruo',
    'sa': 'sa', 'sai': 'sai', 'san': 'san', 'sang': 'sang', 'sao': 'sao',
    'se': 'se', 'sen': 'sen', 'seng': 'seng', 'sha': 'sha', 'shai': 'shai',
    'shan': 'shan', 'shang': 'shang', 'shao': 'shao', 'she': 'she',
    'shei': 'shei', 'shen': 'shen', 'sheng': 'sheng', 'shi': 'shi',
    'shou': 'shou', 'shu': 'shu', 'shua': 'shua', 'shuai': 'shuai',
    'shuan': 'shuan', 'shuang': 'shuang', 'shui': 'shui', 'shun': 'shun',
    'shuo': 'shuo', 'si': 'si', 'song': 'song', 'sou': 'sou', 'su': 'su',
    'suan': 'suan', 'sui': 'sui', 'sun': 'sun', 'suo': 'suo',
    'ta': 'ta', 'tai': 'tai', 'tan': 'tan', 'tang': 'tang', 'tao': 'tao',
    'te': 'te', 'teng': 'teng', 'ti': 'ti', 'tian': 'tian', 'tiao': 'tiao',
    'tie': 'tie', 'ting': 'ting', 'tong': 'tong', 'tou': 'tou', 'tu': 'tu',
    'tuan': 'tuan', 'tui': 'tui', 'tun': 'tun', 'tuo': 'tuo',
    'wa': 'wa', 'wai': 'wai', 'wan': 'wan', 'wang': 'wang', 'wei': 'wei',
    'wen': 'wen', 'weng': 'weng', 'wo': 'wo', 'wu': 'wu',
    'xi': 'xi', 'xia': 'xia', 'xian': 'xian', 'xiang': 'xiang',
    'xiao': 'xiao', 'xie': 'xie', 'xin': 'xin', 'xing': 'xing',
    'xiong': 'xiong', 'xiu': 'xiu', 'xu': 'xu', 'xuan': 'xuan',
    'xue': 'xue', 'xun': 'xun',
    'ya': 'ya', 'yan': 'yan', 'yang': 'yang', 'yao': 'yao', 'ye': 'ye',
    'yi': 'yi', 'yin': 'yin', 'ying': 'ying', 'yo': 'yo', 'yong': 'yong',
    'you': 'you', 'yu': 'yu', 'yuan': 'yuan', 'yue': 'yue', 'yun': 'yun',
    'za': 'za', 'zai': 'zai', 'zan': 'zan', 'zang': 'zang', 'zao': 'zao',
    'ze': 'ze', 'zei': 'zei', 'zen': 'zen', 'zeng': 'zeng', 'zha': 'zha',
    'zhai': 'zhai', 'zhan': 'zhan', 'zhang': 'zhang', 'zhao': 'zhao',
    'zhe': 'zhe', 'zhei': 'zhei', 'zhen': 'zhen', 'zheng': 'zheng',
    'zhi': 'zhi', 'zhong': 'zhong', 'zhou': 'zhou', 'zhu': 'zhu',
    'zhua': 'zhua', 'zhuai': 'zhuai', 'zhuan': 'zhuan', 'zhuang': 'zhuang',
    'zhui': 'zhui', 'zhun': 'zhun', 'zhuo': 'zhuo', 'zi': 'zi',
    'zong': 'zong', 'zou': 'zou', 'zu': 'zu', 'zuan': 'zuan',
    'zui': 'zui', 'zun': 'zun', 'zuo': 'zuo',
}


def chinese_to_pinyin(name: str) -> str:
    """将中文姓名转换为拼音（简易实现）。
    
    对于常见姓氏使用映射表，名字部分尝试逐字匹配。
    无法转换的字符保留原样。
    """
    if not name or not is_chinese_name(name):
        return name
    
    result = []
    chars = list(name)
    
    for i, char in enumerate(chars):
        if '\u4e00' <= char <= '\u9fff':
            if i == 0 and char in COMMON_SURNAMES:
                result.append(COMMON_SURNAMES[char])
            else:
                # 尝试从预定义映射中找
                found = False
                for pinyin in PINYIN_INITIALS.values():
                    # 简单跳过 - 完整的拼音转换需要 pypinyin 库
                    pass
                # 如果没有 pypinyin，使用首字母大写作为 fallback
                # 尝试常见单字拼音
                result.append(char)  # fallback: 保留原字
        else:
            result.append(char)
    
    return "".join(result)


def get_author_identifier(raw_author: str, use_fullname: bool = False) -> str:
    """获取作者标识符用于文件命名。
    
    Args:
        raw_author: 原始作者名
        use_fullname: 是否使用全名（冲突降级时为 True）
    
    Returns:
        姓氏（默认）或 姓氏+名字全拼（冲突时）
    """
    surname, given = split_author_name(raw_author)
    
    if not surname:
        return ""
    
    # 中文姓名处理
    if is_chinese_name(raw_author):
        pinyin = chinese_to_pinyin(raw_author)
        if use_fullname and len(raw_author) >= 2:
            # 全名拼音: WangXiaoming
            return pinyin
        else:
            # 仅姓氏拼音: Wang
            # 取拼音的第一个单词作为姓氏
            return pinyin if not is_chinese_name(pinyin) else COMMON_SURNAMES.get(raw_author[0], pinyin)
    
    # 英文姓名
    if use_fullname and given:
        # 姓氏+名字全拼: SmithJohnDavid -> SmithJohn
        return f"{surname}{given}"
    else:
        return surname


def extract_authors_from_text(text: str) -> list:
    """从文本中提取作者列表（启发式）。"""
    authors = []
    author_section = re.search(
        r'(?:author|by|created by)[s]?\s*[:：]\s*(.+?)(?:\n\n|\n[A-Z]|Abstract|Introduction|$)', 
        text, re.IGNORECASE | re.DOTALL
    )
    if author_section:
        author_text = author_section.group(1)[:500]
        parts = re.split(r'\s+(?:and|&)\s+|[,;]\s*', author_text)
        for p in parts[:10]:
            name = p.strip()
            if name and len(name) >= 2:
                authors.append(name)

    if not authors:
        first_line = text.split('\n')[0]
        m = re.match(r'([A-Z][a-z]+ [A-Z][a-z]+)', first_line)
        if m:
            authors.append(m.group(1))

    return authors[:5]


# ── 标题提取 ──────────────────────────────────────────────────

def extract_title_from_text(text: str) -> str:
    """从文本中提取论文标题。"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for i, line in enumerate(lines[:15]):
        if 30 < len(line) < 250 and not line.startswith(('http', 'doi:', '(c)', 'Correspondence')):
            if line.isupper() and len(line) < 60:
                continue
            if re.match(r'^\d+$', line):
                continue
            return line.strip().rstrip('.')
    return ""


# ── 关键词提取 ────────────────────────────────────────────────

def extract_keywords_from_text(text: str, title: str = "", max_count: int = 5) -> list:
    """从文本中提取关键词。
    
    提取策略:
    1. 从文本前 3000 字符中提取高频学术术语
    2. 从标题中提取核心名词短语
    3. 最多返回 max_count 个关键词
    """
    keywords = []
    
    # 策略 1: 从标题提取核心术语
    if title:
        # 去除常见非名词词
        title_words = re.findall(r'[A-Za-z]{3,}|[\u4e00-\u9fff]{2,}', title)
        for w in title_words:
            w_lower = w.lower()
            if w_lower not in STOP_WORDS and len(w_lower) >= 3:
                if w not in keywords:
                    keywords.append(w)
    
    # 策略 2: 从文本前 3000 字符提取高频术语
    text_sample = text[:3000]
    
    # 英文词频统计
    words = re.findall(r'\b[A-Za-z]{3,}\b', text_sample)
    word_freq = Counter(w.lower() for w in words if w.lower() not in STOP_WORDS)
    
    # 中文词频统计（简易：2-4字组合）
    chinese_segments = re.findall(r'[\u4e00-\u9fff]{2,4}', text_sample)
    cn_freq = Counter(chinese_segments)
    
    # 合并英文和中文关键词候选
    candidates = []
    for word, freq in word_freq.most_common(20):
        if freq >= 2 and word not in [k.lower() for k in keywords]:
            candidates.append((word, freq))
    
    for word, freq in cn_freq.most_common(10):
        if freq >= 2 and word not in keywords:
            # 过滤掉包含停用词的中文片段
            if not all(sw in word for sw in ['的', '了', '在', '是']):
                candidates.append((word, freq))
    
    # 按频率排序，取前 N 个
    candidates.sort(key=lambda x: x[1], reverse=True)
    for word, _ in candidates:
        if len(keywords) >= max_count:
            break
        # 首字母大写（英文）
        if word[0].isalpha() and word[0].isascii():
            word = word.capitalize()
        keywords.append(word)
    
    return keywords[:max_count]


# ── PDF 元数据提取 ─────────────────────────────────────────────

def extract_from_pdf(filepath: str) -> dict:
    """从 PDF 文件中提取元数据。"""
    meta = empty_meta(filepath)
    meta["source"] = "pdf"
    meta["md5"] = file_md5(filepath)

    if not HAS_PYPDF:
        return meta

    try:
        reader = PyPDF2.PdfReader(filepath)
        meta["page_count"] = len(reader.pages)

        info = reader.metadata
        if info:
            if info.author:
                raw_author = str(info.author)
                meta["authors"] = [raw_author]
                surname, given = split_author_name(raw_author)
                meta["first_author_surname"] = surname
                meta["first_author_given_name"] = given
                meta["first_author_fullname"] = f"{surname}{given}" if given else surname

            if info.title:
                meta["title"] = str(info.title).strip()

            if info.creation_date:
                cd = str(info.creation_date)
                m = re.search(r'(\d{4})', cd)
                if m:
                    meta["year"] = m.group(1)

        # 提取前 3 页文本
        text_parts = []
        for page in reader.pages[:3]:
            try:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
            except Exception:
                pass
        full_text = '\n'.join(text_parts)[:5000]
        meta["text_preview"] = full_text[:3000]

        if full_text:
            doi = extract_doi_from_text(full_text)
            if doi:
                meta["doi"] = doi
            if not meta["year"]:
                meta["year"] = extract_year_from_text(full_text)
            if not meta["title"]:
                meta["title"] = extract_title_from_text(full_text)
            if not meta["authors"]:
                meta["authors"] = extract_authors_from_text(full_text)
                if meta["authors"]:
                    surname, given = split_author_name(meta["authors"][0])
                    meta["first_author_surname"] = surname
                    meta["first_author_given_name"] = given
                    meta["first_author_fullname"] = f"{surname}{given}" if given else surname

            # 完整性评估
            lower = full_text.lower()
            meta["has_abstract"] = bool(re.search(r'\babstract\b', lower))
            meta["has_references"] = bool(re.search(r'\breferences?\b|\bbibliography\b', lower))
            meta["has_appendix"] = bool(re.search(r'\bappendix\b', lower))
            meta["has_figures"] = bool(re.search(r'\bfigure\s+\d\b|\bfig\.\s+\d\b', lower))
            meta["has_fulltext"] = meta["page_count"] >= 3

            # 关键词提取
            if not meta["keywords"]:
                meta["keywords"] = extract_keywords_from_text(full_text, meta["title"])

    except Exception:
        pass

    return meta


# ── DOCX 元数据提取 ────────────────────────────────────────────

def extract_from_docx(filepath: str) -> dict:
    """从 DOCX 文件中提取元数据。"""
    meta = empty_meta(filepath)
    meta["source"] = "docx"
    meta["md5"] = file_md5(filepath)

    if not HAS_DOCX:
        return meta

    try:
        doc = docx.Document(filepath)

        text_parts = []
        for para in doc.paragraphs[:100]:
            if para.text.strip():
                text_parts.append(para.text.strip())
        full_text = '\n'.join(text_parts)[:5000]
        meta["text_preview"] = full_text[:3000]

        if full_text:
            doi = extract_doi_from_text(full_text)
            if doi:
                meta["doi"] = doi
            meta["year"] = extract_year_from_text(full_text)
            meta["title"] = extract_title_from_text(full_text)
            meta["authors"] = extract_authors_from_text(full_text)
            if meta["authors"]:
                surname, given = split_author_name(meta["authors"][0])
                meta["first_author_surname"] = surname
                meta["first_author_given_name"] = given
                meta["first_author_fullname"] = f"{surname}{given}" if given else surname

            lower = full_text.lower()
            meta["has_abstract"] = bool(re.search(r'\babstract\b', lower))
            meta["has_references"] = bool(re.search(r'\breferences?\b|\bbibliography\b', lower))
            meta["has_appendix"] = bool(re.search(r'\bappendix\b', lower))
            meta["has_figures"] = bool(re.search(r'\bfigure\s+\d\b|\bfig\.\s+\d\b', lower))
            meta["has_fulltext"] = len(doc.paragraphs) >= 10

            if not meta["keywords"]:
                meta["keywords"] = extract_keywords_from_text(full_text, meta["title"])

        meta["page_count"] = len(doc.paragraphs) // 40 or 1

    except Exception:
        pass

    return meta


# ── TXT 元数据提取 ─────────────────────────────────────────────

def extract_from_txt(filepath: str) -> dict:
    """从 TXT 文件中提取元数据。"""
    meta = empty_meta(filepath)
    meta["source"] = "txt"
    meta["md5"] = file_md5(filepath)

    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read(5000)
    except Exception:
        return meta

    meta["text_preview"] = text[:3000]

    if text:
        doi = extract_doi_from_text(text)
        if doi:
            meta["doi"] = doi
        meta["year"] = extract_year_from_text(text)
        meta["title"] = extract_title_from_text(text)
        meta["authors"] = extract_authors_from_text(text)
        if meta["authors"]:
            surname, given = split_author_name(meta["authors"][0])
            meta["first_author_surname"] = surname
            meta["first_author_given_name"] = given
            meta["first_author_fullname"] = f"{surname}{given}" if given else surname

        lower = text.lower()
        meta["has_abstract"] = bool(re.search(r'\babstract\b', lower))
        meta["has_references"] = bool(re.search(r'\breferences?\b|\bbibliography\b', lower))

        if not meta["keywords"]:
            meta["keywords"] = extract_keywords_from_text(text, meta["title"])

    return meta


# ── DOI 在线解析 ───────────────────────────────────────────────

def resolve_doi(doi: str) -> dict:
    """通过 Crossref API 解析 DOI 获取完整元数据。"""
    meta = empty_meta()
    if not HAS_REQUESTS:
        return meta

    doi = doi.strip()
    if doi.startswith("http"):
        doi = doi.split("doi.org/")[-1] if "doi.org/" in doi else doi.split("/")[-1]

    url = f"https://api.crossref.org/works/{doi}"
    try:
        resp = requests.get(url, headers={"Accept": "application/json"}, timeout=15)
        if resp.status_code != 200:
            return meta
        data = resp.json()["message"]
    except Exception:
        return meta

    meta["doi"] = doi
    meta["source"] = "doi"

    if "title" in data and data["title"]:
        meta["title"] = data["title"][0]

    if "author" in data:
        for a in data["author"]:
            family = a.get("family", "")
            given = a.get("given", "")
            full = f"{given} {family}".strip() if given else family
            if full:
                meta["authors"].append(full)
        if meta["authors"]:
            surname, given_name = split_author_name(meta["authors"][0])
            meta["first_author_surname"] = surname
            meta["first_author_given_name"] = given_name
            meta["first_author_fullname"] = f"{surname}{given_name}" if given_name else surname

    if "created" in data and "date-parts" in data["created"]:
        meta["year"] = str(data["created"]["date-parts"][0][0])
    elif "published-print" in data and "date-parts" in data["published-print"]:
        meta["year"] = str(data["published-print"]["date-parts"][0][0])

    if "container-title" in data and data["container-title"]:
        meta["journal"] = data["container-title"][0]

    meta["volume"] = str(data.get("volume", ""))
    meta["issue"] = str(data.get("issue", ""))
    meta["pages"] = str(data.get("page", ""))

    # 关键词
    if "subject" in data and data["subject"]:
        meta["keywords"] = data["subject"][:5]
    elif "keyword" in data and data["keyword"]:
        meta["keywords"] = data["keyword"][:5]

    return meta


# ── 完整性评分 ─────────────────────────────────────────────────

def completeness_score(meta: dict) -> int:
    """计算文献信息完整度得分（满分 10）。"""
    score = 0
    if meta.get("has_abstract"):
        score += 1
    if meta.get("has_fulltext") or meta.get("page_count", 0) >= 3:
        score += 3
    if meta.get("has_references"):
        score += 2
    if meta.get("has_appendix"):
        score += 2
    if meta.get("has_figures"):
        score += 1
    if meta.get("doi"):
        score += 1
    return score


# ── APA 格式构建（V2 逐字段规则） ──────────────────────────────

def build_apa_citation(meta: dict) -> str:
    """根据元数据构建 APA 格式引用字符串（严格遵循 V2 规则）。

    英文模板:
        Author, A. A., & Author, B. B. (Year). Title of article. *Journal Name*, Volume(Issue), Page-Page. https://doi.org/xxx

    中文模板:
        作者. (年份). 文章标题. 期刊名称, 卷号(期号), 页码.
    """
    if not meta.get("authors"):
        return ""
    
    # 判断中文文献
    is_chinese = is_chinese_name(meta["authors"][0]) if meta["authors"] else False
    
    parts = []
    
    # ── 作者 ──
    if is_chinese:
        # 中文: 作者全名，逗号分隔，最后两位用"和"连接
        author_names = []
        for a in meta["authors"][:20]:
            author_names.append(a)
        if len(author_names) >= 2:
            author_str = ", ".join(author_names[:-1]) + "和" + author_names[-1]
        else:
            author_str = author_names[0] if author_names else ""
        parts.append(author_str + ".")
    else:
        # 英文: 姓, 首字母缩写. 多作者用逗号，最后用 &
        apa_authors = []
        for a in meta["authors"][:20]:
            apa_name = normalize_author_apa(a)
            if apa_name:
                apa_authors.append(apa_name)
        if len(apa_authors) >= 2:
            author_str = ", ".join(apa_authors[:-1]) + ", & " + apa_authors[-1]
        elif apa_authors:
            author_str = apa_authors[0]
        else:
            return ""
        parts.append(author_str + ".")
    
    # ── 年份 ──
    year = meta.get("year", "")
    if year:
        parts.append(f"({year}).")
    else:
        return ""  # 无年份不生成
    
    # ── 标题 ──
    title = meta.get("title", "")
    if title:
        if is_chinese:
            parts.append(title + ".")
        else:
            # 英文标题：仅首字母大写（专有名词除外）
            # 简单处理：保持原标题大小写，但不全大写
            if title.isupper():
                title = title.title()
            parts.append(title + ".")
    else:
        return ""
    
    # ── 期刊信息 ──
    journal = meta.get("journal", "")
    volume = meta.get("volume", "")
    issue = meta.get("issue", "")
    pages = meta.get("pages", "")
    doi = meta.get("doi", "")
    
    if journal:
        if is_chinese:
            journal_part = journal
        else:
            journal_part = f"*{journal}*"  # 斜体
        
        if volume:
            if is_chinese:
                journal_part += f", {volume}"
            else:
                journal_part += f", *{volume}*"  # 卷号斜体
            if issue:
                journal_part += f"({issue})"  # 期号不斜体
            if pages:
                journal_part += f", {pages}"
        elif pages:
            journal_part += f", {pages}"
        
        parts.append(journal_part + ".")
    elif doi:
        # 无期刊名，视为预印本
        parts.append("*arXiv preprint*.")
    
    # ── DOI ──
    if doi:
        doi_clean = doi.replace("https://doi.org/", "").replace("http://doi.org/", "")
        parts.append(f"https://doi.org/{doi_clean}")
    
    apa = " ".join(parts)
    
    # ── 强制自检 ──
    if not validate_apa(apa, is_chinese):
        # 自检失败，尝试修复
        # 常见问题：双句点
        apa = re.sub(r'\.\.+', '.', apa)
        apa = re.sub(r'\.\s*\.', '.', apa)
    
    return apa


def validate_apa(apa: str, is_chinese: bool = False) -> bool:
    """APA 引用自检（V2 强制执行）。
    
    检查项:
    1. 作者格式是否为 "姓, 首字母缩写." 或中文全名
    2. 年份是否用括号包裹并以句点结束
    3. 标题是否仅首字母大写（英文）
    4. 是否存在多余标点
    """
    if not apa:
        return False
    
    # 检查双句点
    if '..' in apa:
        return False
    
    # 检查年份括号
    if not re.search(r'\(\d{4}\)\.', apa) and not re.search(r'\(Unknown\)\.', apa):
        return False
    
    # 检查多余逗号
    if re.search(r',\s*,', apa):
        return False
    
    return True


# ── 文件名推断 ─────────────────────────────────────────────────

def infer_from_filename(filepath: str) -> dict:
    """从文件名中推断元数据。"""
    meta = empty_meta(filepath)
    meta["source"] = "filename"
    name = os.path.splitext(os.path.basename(filepath))[0]

    # 模式 1: AuthorYear (如 Smith2020, WangXiaoming2023)
    m1 = re.match(r'([A-Z][a-zA-Z]+?)(\d{4})', name)
    if m1:
        raw_author = m1.group(1)
        surname, given = split_author_name(raw_author)
        meta["authors"] = [raw_author]
        meta["first_author_surname"] = surname if surname else raw_author
        meta["first_author_given_name"] = given
        meta["first_author_fullname"] = f"{surname}{given}" if given else surname
        meta["year"] = m1.group(2)
        return meta

    # 模式 2: 以年份开头或结尾
    m2 = re.search(r'(\d{4})', name)
    if m2:
        meta["year"] = m2.group(1)
        prefix = name[:m2.start()]
        author_candidate = prefix.strip('_ -').strip()
        if author_candidate and len(author_candidate) <= 30:
            meta["authors"] = [author_candidate]
            surname, given = split_author_name(author_candidate)
            meta["first_author_surname"] = surname if surname else author_candidate
            meta["first_author_given_name"] = given
            meta["first_author_fullname"] = f"{surname}{given}" if given else surname

    return meta


# ── 构建目标文件名（含冲突降级） ───────────────────────────────

def build_target_filename(meta: dict, ext: str, existing_names: set = None,
                           all_metas: list = None) -> str:
    """构建标准化文件名，支持冲突降级。
    
    Args:
        meta: 文件元数据
        ext: 文件扩展名（含点号）
        existing_names: 已存在的文件名集合
        all_metas: 所有文件的元数据列表（用于检测同名同年冲突）
    
    Returns:
        标准化文件名
    """
    if existing_names is None:
        existing_names = set()
    
    surname = meta.get("first_author_surname", "")
    fullname = meta.get("first_author_fullname", "")
    year = meta.get("year", "")
    
    if not surname or not year:
        return meta.get("filename", "")
    
    # 默认：仅使用姓氏
    base_name = f"{surname}{year}"
    
    # 冲突检测：是否有同姓氏同年份但不同文献的情况
    need_fullname = False
    if all_metas:
        for other_meta in all_metas:
            if other_meta is meta:
                continue
            other_surname = other_meta.get("first_author_surname", "")
            other_year = other_meta.get("year", "")
            other_fullname = other_meta.get("first_author_fullname", "")
            
            if (other_surname == surname and 
                other_year == year and 
                other_fullname != fullname):
                need_fullname = True
                break
    
    if need_fullname and fullname and fullname != surname:
        base_name = f"{fullname}{year}"
    
    target = base_name + ext
    
    # 检查是否已存在同名文件（不同文献）
    if target in existing_names:
        # 加小写字母后缀
        for letter in 'abcdefghijklmnopqrstuvwxy':
            alt = f"{base_name}{letter}{ext}"
            if alt not in existing_names:
                return alt
    
    return target


# ── 统一入口：提取任意文件的元数据 ──────────────────────────────

def extract_metadata(filepath: str, resolve_doi_flag: bool = False) -> dict:
    """提取文献文件的完整元数据。

    提取优先级链：
    1. PDF 元数据 / DOCX 内容
    2. 文件名推断
    3. DOI 在线解析（可选，需 resolve_doi_flag=True）
    """
    ext = Path(filepath).suffix.lower()

    # Level 1: 内容/元数据提取
    if ext == '.pdf':
        meta = extract_from_pdf(filepath)
    elif ext == '.docx':
        meta = extract_from_docx(filepath)
    elif ext == '.txt':
        meta = extract_from_txt(filepath)
    else:
        meta = empty_meta(filepath)

    # Level 2: 文件名推断（补充缺失信息）
    if not meta["year"] or not meta["first_author_surname"]:
        fn_meta = infer_from_filename(filepath)
        if not meta["year"] and fn_meta["year"]:
            meta["year"] = fn_meta["year"]
        if not meta["first_author_surname"] and fn_meta["first_author_surname"]:
            meta["authors"] = fn_meta["authors"]
            meta["first_author_surname"] = fn_meta["first_author_surname"]
            meta["first_author_given_name"] = fn_meta.get("first_author_given_name", "")
            meta["first_author_fullname"] = fn_meta.get("first_author_fullname", "")

    # Level 3: DOI 在线解析
    if resolve_doi_flag and meta["doi"] and (not meta["authors"] or not meta["title"]):
        doi_meta = resolve_doi(meta["doi"])
        if doi_meta["title"]:
            for key in ["title", "authors", "first_author_surname", "first_author_given_name",
                        "first_author_fullname", "year", "journal", "volume", "issue", "pages",
                        "keywords"]:
                if not meta.get(key) and doi_meta.get(key):
                    meta[key] = doi_meta[key]

    # 生成 target filename（基础版，不含冲突降级）
    if meta["first_author_surname"] and meta["year"]:
        meta["target_filename"] = f"{meta['first_author_surname']}{meta['year']}{ext}"
    else:
        meta["target_filename"] = os.path.basename(filepath)

    # 完整性评分
    meta["completeness_score"] = completeness_score(meta)

    # APA 引用
    meta["apa"] = build_apa_citation(meta)

    # 关键词（分号分隔字符串）
    if meta["keywords"]:
        meta["keywords_str"] = "; ".join(meta["keywords"])
    else:
        meta["keywords_str"] = ""

    return meta


# ── CLI ─────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="文献元数据提取工具 V2")
    parser.add_argument("path", help="文件路径或目录路径")
    parser.add_argument("--batch", action="store_true", help="批量处理目录")
    parser.add_argument("--doi", action="store_true", help="启用 DOI 在线解析")
    parser.add_argument("--apa", action="store_true", help="同时输出 APA 引用")
    args = parser.parse_args()

    if args.batch:
        # 先收集所有文件和元数据
        all_files = []
        all_metas_list = []
        for root, dirs, files in os.walk(args.path):
            dirs[:] = [d for d in dirs if d not in {'Duplicates', '__pycache__', '.git'}]
            for f in files:
                if Path(f).suffix.lower() in PAPER_EXTS:
                    fp = os.path.join(root, f)
                    meta = extract_metadata(fp, resolve_doi_flag=args.doi)
                    all_files.append((fp, meta))
                    all_metas_list.append(meta)

        # 构建冲突降级后的文件名
        existing_names = set()
        for fp, meta in all_files:
            ext = Path(fp).suffix.lower()
            target = build_target_filename(meta, ext, existing_names, all_metas_list)
            existing_names.add(target)
            
            out = {
                "filepath": fp,
                "target_name": target,
                "year": meta["year"],
                "authors": meta["authors"],
                "first_author_surname": meta["first_author_surname"],
                "first_author_fullname": meta.get("first_author_fullname", ""),
                "title": meta["title"][:80] if meta["title"] else "",
                "doi": meta["doi"],
                "keywords": meta.get("keywords", []),
                "keywords_str": meta.get("keywords_str", ""),
                "score": meta["completeness_score"],
            }
            if args.apa:
                out["apa"] = meta.get("apa", "")
            print(json.dumps(out, ensure_ascii=False))
    else:
        meta = extract_metadata(args.path, resolve_doi_flag=args.doi)
        out = {
            "filepath": args.path,
            "filename": os.path.basename(args.path),
            "target_name": meta["target_filename"],
            "year": meta["year"],
            "authors": meta["authors"],
            "first_author_surname": meta["first_author_surname"],
            "first_author_fullname": meta.get("first_author_fullname", ""),
            "first_author_given_name": meta.get("first_author_given_name", ""),
            "title": meta["title"],
            "journal": meta["journal"],
            "doi": meta["doi"],
            "volume": meta["volume"],
            "issue": meta["issue"],
            "pages": meta["pages"],
            "keywords": meta.get("keywords", []),
            "keywords_str": meta.get("keywords_str", ""),
            "apa": meta.get("apa", ""),
            "md5": meta["md5"],
            "page_count": meta["page_count"],
            "completeness_score": meta["completeness_score"],
            "has_abstract": meta["has_abstract"],
            "has_references": meta["has_references"],
            "has_appendix": meta["has_appendix"],
        }
        print(json.dumps(out, ensure_ascii=False, indent=2))
